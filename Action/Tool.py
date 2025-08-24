# tool_definitions.py
import os
import uuid
from dotenv import load_dotenv
from langchain_core.tools import tool
from tavily import TavilyClient
from Call_llm import LLM_Pipeline
import docx
from datetime import datetime
# Load variables from .env file
load_dotenv()



@tool
def read_google_sheet(file_name: str) -> list[dict]:
    """
    Reads data from a Google Sheet and returns it as a list of dictionaries.

    Args:
      file_name: The name of the Google Sheet to read.

    Returns:
      A list of dictionaries representing the rows of the sheet.
    """
    print(f"--- Calling Google Sheet Tool: Reading '{file_name}' ---")

    # Mock data representing the sheet's content for this example
    mock_sheets_db = {
        "Q3 Sales": [
    {"company": "Innovate Inc.", "sales": 995000000},
    {"company": "QuantumLeap Co.", "sales": 7000000},
    {"company": "DataWeavers", "sales": 3500000},
    {"company": "AlphaTech", "sales": 650000},
    {"company": "KPMG", "sales": 38400000000},
    {"company": "Infosys", "sales": 19280000000},
    {"company": "NextGen Partners", "sales": 7600000}
]}
    if file_name in mock_sheets_db:
        return mock_sheets_db[file_name]
    else:
        # In a real agent, this would raise an error to be caught
        # by the Executor for a recovery attempt.
        raise FileNotFoundError(f"The sheet '{file_name}' was not found.")

@tool
def get_highest_sale_record(data : list[dict]) -> dict:
    """
    Use this tool to analyze a list of sales records and identify the single record 
    with the highest sales value. The input should be a list of dictionaries, 
    where each dictionary represents a company's sales data.

    Each dictionary in the list MUST contain a 'company' key with a string value 
    and a 'sales' key with a numeric value.
    """
    print("--- Finding the record with the highest sales ---")
    
    if not data:
        return {"error": "No data provided."}
    
    # Use max() to find the record with the highest 'sales' value
    highest_sales_record = max(data, key=lambda record: record['sales'])
    
    return highest_sales_record

@tool
def search_recent_company_news(company_name: str) -> str:
    """
    Searches the web for recent, relevant news articles about a specific 
    company using the Tavily Search API.
    """
    print(f"--- Calling Tavily Search Tool for '{company_name}' news ---")
    
    # 1. Initialize the Tavily Client
    # It automatically uses the TAVILY_API_KEY from your environment variables.
    try:
        tavily_client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))
    except KeyError:
        return ["Error: TAVILY_API_KEY environment variable not set."]

    # 2. Perform the search
    # We construct a query focused on recent news.
    search_query = f"latest financial news and developments for {company_name}"
    
    try:
        response = tavily_client.search(
            query=search_query,
            search_depth="basic", # Use "basic" for speed or "advanced" for more detail
            max_results=10      # Limit to the top 3 most relevant articles
        )
        print(f"Raw response {response}")
        extracted_content = [result['content'] for result in response['results']]
        content = "".join(extracted_content)
        print(f"Extracted content: {content}")
        return content
    except Exception as e:
        return [f"Error during Tavily search: {e}"]



@tool
def create_google_doc(title: str, content: str) -> str:
    """
    Use this tool to create a new Google Doc with a specific title and content. 
    This is the final step for presenting a summary or report.

    Args:
        title: The title for the new Google Doc.
        content: The text content to be written into the document.

    Returns:
        A confirmation message with the new document's URL.
    """
    print(f"--- TOOL: Simulating creation of Google Doc titled '{title}' ---")
    
    # Generate a unique-looking ID for the fake document
    doc_id = str(uuid.uuid4())
    output_dir = "../output_documents"
    os.makedirs(output_dir, exist_ok=True)
    # 2. Create a safe filename base from the title
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '_')).rstrip()
    safe_title = safe_title.replace(' ', '_')
     # NEW: Generate a unique timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # NEW: Combine the title and timestamp for the final filename
    final_filename = f"{safe_title}_{timestamp}.docx"
    file_path = os.path.join(output_dir, final_filename)
    document = docx.Document()
    document.add_heading(title, level=1)
    document.add_paragraph(content)
    document.save(file_path)
    # In a real implementation, this is where you would use the Google Docs API.
    # For now, we just print the content that would be written.
    print("\n--- Document Content ---")
    print(content)
    print("------------------------")
    
    # fake_url = f"https://docs.google.com/document/d/{doc_id}"
    fake_url = file_path
    return f"Successfully created Google Doc titled '{title}'. You can view it here: {fake_url}"


# The agent can now directly use these functions.
# LangChain handles the registry creation in the background.
tools = [read_google_sheet, get_highest_sale_record, search_recent_company_news, create_google_doc]

if __name__ == "__main__":
  
    data = read_google_sheet.invoke("Q3 Sales")
    print(f"Data from Google Sheet: {data}")
    # print(type(data['Q3 Sales']))
    # input_data = {"data": mock_db["Q3 Sales"]}
    higest = get_highest_sale_record.invoke({"data": data})
    print(f"Record with the highest sales: ----------------------\n {higest}")
# 
    # news = search_recent_company_news("KPMG")
    # print(f"Recent news about kpmg: \n---------------------------\n{news}")

